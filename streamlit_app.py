import streamlit as st
import logging
import io
import re
import os
import pandas as pd
import tempfile
import subprocess
import base64

from pdf2image import convert_from_bytes
import pytesseract
from pytesseract import Output
from docx import Document
from openpyxl.cell.cell import ILLEGAL_CHARACTERS_RE
from PIL import Image, ImageDraw
from llama_index import RecursiveCharacterTextSplitter  # Single import

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')


def convert_pdf_to_images(pdf_file: io.BytesIO, poppler_path: str = None) -> list:
    """
    Converts a PDF file (as a BytesIO stream) into a list of image objects.

    Args:
        pdf_file (io.BytesIO): The input PDF file stream.
        poppler_path (str, optional): Path to the Poppler binaries (if required). 
            Example (Windows): r'C:\poppler\bin'.

    Returns:
        list: A list of PIL.Image objects representing each PDF page.
            Example: [Image object (Page 1), Image object (Page 2), ...]
    """
    try:
        pdf_file.seek(0)
        pdf_bytes = pdf_file.read()
        images = convert_from_bytes(pdf_bytes, poppler_path=poppler_path)
        logging.info("Converted PDF to %d image(s).", len(images))
        return images
    except Exception as e:
        logging.error("Error converting PDF to images: %s", e)
        st.error("Failed to process PDF file. Ensure Poppler is installed and in PATH.")
        return []


def extract_text_from_images(images: list) -> list:
    """
    Extracts full-page text from a list of image objects using Tesseract OCR.

    Args:
        images (list): List of PIL.Image objects.
            Example: [Image object (Page 1), Image object (Page 2), ...]

    Returns:
        list: List of text strings extracted from each image.
            Example: ["Text from page 1", "Text from page 2", ...]
    """
    texts = []
    for idx, image in enumerate(images, start=1):
        try:
            text = pytesseract.image_to_string(image)
            texts.append(text)
            logging.info("Extracted text from page %d.", idx)
        except Exception as e:
            logging.error("Error extracting text from page %d: %s", idx, e)
            texts.append("")
    return texts


def create_excel_file(texts: list) -> io.BytesIO:
    """
    Creates an Excel file from a list of text strings after cleaning illegal characters.

    Args:
        texts (list): List of text strings.
            Example: ["Text for page 1", "Text for page 2", ...]

    Returns:
        io.BytesIO: A BytesIO stream containing the Excel file.
            Example: BytesIO stream for 'extracted_text.xlsx'
    """
    def clean_text(text: str) -> str:
        return re.sub(ILLEGAL_CHARACTERS_RE, "", text) if text else text

    cleaned_texts = [clean_text(text) for text in texts]
    df = pd.DataFrame({
        'Page': list(range(1, len(cleaned_texts) + 1)),
        'Content': cleaned_texts
    })
    output = io.BytesIO()
    try:
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='ExtractedText')
        output.seek(0)
        logging.info("Excel file created with %d page(s).", len(cleaned_texts))
    except Exception as e:
        logging.error("Error creating Excel file: %s", e)
        st.error("Failed to create Excel file.")
    return output


def create_word_file(texts: list) -> io.BytesIO:
    """
    Creates a Word document from a list of text strings.

    Args:
        texts (list): List of text strings.
            Example: ["Text for page 1", "Text for page 2", ...]

    Returns:
        io.BytesIO: A BytesIO stream containing the Word document.
            Example: BytesIO stream for 'extracted_text.docx'
    """
    document = Document()
    for idx, text in enumerate(texts, start=1):
        try:
            document.add_heading(f'Page {idx}', level=1)
            document.add_paragraph(text)
            logging.info("Added page %d to Word document.", idx)
        except Exception as e:
            logging.error("Error adding page %d to Word document: %s", idx, e)
    output = io.BytesIO()
    try:
        document.save(output)
        output.seek(0)
        logging.info("Word document created with %d page(s).", len(texts))
    except Exception as e:
        logging.error("Error saving Word document: %s", e)
        st.error("Failed to create Word document.")
    return output


def extract_tables_from_images_with_tesseract(images: list) -> list:
    """
    Extracts table-like data from images using Tesseract's TSV output.
    For each image, returns a pandas DataFrame of the TSV data.

    Args:
        images (list): List of PIL.Image objects.
            Example: [Image object (Page 1), Image object (Page 2), ...]

    Returns:
        list: A list of pandas DataFrame objects containing TSV data.
            Example: [DataFrame for page 1, DataFrame for page 2, ...]
    """
    tables = []
    for idx, image in enumerate(images, start=1):
        try:
            df = pytesseract.image_to_data(image, output_type=Output.DATAFRAME)
            df = df[df['text'].notna() & (df['text'] != "")]
            logging.info("Extracted TSV data for page %d with %d rows.", idx, len(df))
            tables.append(df)
        except Exception as e:
            logging.error("Error extracting TSV data from page %d: %s", idx, e)
    return tables


def draw_bounding_boxes_on_image(image: Image.Image, conf_threshold: int = 60) -> Image.Image:
    """
    Draws red bounding boxes on an image around detected text regions (using Tesseract).
    Only boxes with confidence above the threshold are drawn.

    Args:
        image (PIL.Image): The input image.
            Example: PIL.Image.open("page1.png")
        conf_threshold (int, optional): Confidence threshold. Defaults to 60.
            Example: 60

    Returns:
        PIL.Image: The image with red bounding boxes.
            Example: Modified image with boxes around text.
    """
    draw = ImageDraw.Draw(image)
    data = pytesseract.image_to_data(image, output_type=Output.DICT)
    n_boxes = len(data['level'])
    for i in range(n_boxes):
        try:
            conf = int(data['conf'][i])
        except (ValueError, TypeError):
            conf = 0
        if conf > conf_threshold:
            x, y, w, h = data['left'][i], data['top'][i], data['width'][i], data['height'][i]
            draw.rectangle([x, y, x + w, y + h], outline="red", width=2)
    return image


def process_pdf_with_ocrmypdf(pdf_file: io.BytesIO) -> bytes:
    """
    Processes the PDF with OCRmyPDF to add an OCR text layer (making it searchable).
    Uses subprocess to call the 'ocrmypdf' command with the --force-ocr flag.
    Cleans up temporary files after processing.

    Args:
        pdf_file (io.BytesIO): The input PDF file stream.
            Example: io.BytesIO(open("scanned.pdf", "rb").read())

    Returns:
        bytes: The processed PDF as bytes, or None if processing fails.
            Example: b'%PDF-1.4 ...'
    """
    in_path = None
    out_path = None
    try:
        # Write the uploaded PDF to a temporary input file.
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as temp_in:
            temp_in.write(pdf_file.read())
            in_path = temp_in.name
        logging.info("Temporary input file created: %s", in_path)

        # Create a temporary output file.
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as temp_out:
            out_path = temp_out.name
        logging.info("Temporary output file created: %s", out_path)

        # Run OCRmyPDF with --force-ocr to force OCR on every page.
        result = subprocess.run(
            ["ocrmypdf", "--force-ocr", in_path, out_path],
            capture_output=True, text=True, check=True
        )
        logging.info("OCRmyPDF output: %s", result.stdout)

        # Read the processed PDF.
        with open(out_path, "rb") as f:
            processed_pdf = f.read()
        return processed_pdf
    except subprocess.CalledProcessError as e:
        logging.error("OCRmyPDF failed with error: %s", e.stderr)
        st.error("OCRmyPDF processing failed: " + e.stderr)
        return None
    except Exception as e:
        logging.error("Error running OCRmyPDF: %s", e)
        st.error("Failed to process PDF with OCRmyPDF.")
        return None
    finally:
        # Cleanup temporary files
        if in_path and os.path.exists(in_path):
            os.remove(in_path)
            logging.info("Temporary input file removed: %s", in_path)
        if out_path and os.path.exists(out_path):
            os.remove(out_path)
            logging.info("Temporary output file removed: %s", out_path)


def get_semantic_chunks(text: str, chunk_size: int = 512, chunk_overlap: int = 50) -> list:
    """
    Splits the input text into semantic chunks using LlamaIndex's RecursiveCharacterTextSplitter.
    Returns a list of text chunks.

    Args:
        text (str): The input text.
            Example: "This is a long text from OCR..."
        chunk_size (int, optional): Maximum size of each chunk. Defaults to 512.
            Example: 512
        chunk_overlap (int, optional): Overlap between consecutive chunks. Defaults to 50.
            Example: 50

    Returns:
        list: A list of text chunks.
            Example: ["Chunk 1 text...", "Chunk 2 text...", ...]
    """
    try:
        splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        chunks = splitter.split_text(text)
        return chunks
    except Exception as e:
        st.error("Error during semantic chunking: " + str(e))
        return []


# --- Streamlit User Interface ---

st.title("PDF Converter & Table Extractor")
st.write(
    "Upload a scanned PDF. Three output tabs are available:\n\n"
    "1. **Detected Content** – Shows previews with red bounding boxes and export options.\n"
    "2. **OCRmyPDF Output** – Provides a searchable PDF produced by OCRmyPDF.\n"
    "3. **Semantic Chunks** – Displays semantic chunks (based on token/character count) parsed from the OCR text using LlamaIndex."
)

# File uploader (PDF only)
uploaded_pdf = st.file_uploader("Upload a PDF file", type=["pdf"])

# Optional: Specify the Poppler path if needed (e.g., on Windows, set to r'C:\poppler\bin').
poppler_path = None

if uploaded_pdf is not None:
    # Read file bytes once and create separate BytesIO objects for different processing.
    file_bytes = uploaded_pdf.getvalue()
    pdf_for_images = io.BytesIO(file_bytes)
    pdf_for_ocr = io.BytesIO(file_bytes)

    with st.spinner("Converting PDF pages to images..."):
        images = convert_pdf_to_images(pdf_for_images, poppler_path=poppler_path)
    if not images:
        st.error("No images found in the PDF file. Check if the PDF is valid.")
    else:
        # Process the PDF with OCRmyPDF.
        with st.spinner("Running OCRmyPDF on the uploaded PDF..."):
            ocr_pdf = process_pdf_with_ocrmypdf(pdf_for_ocr)

        # Extract OCR text from images.
        with st.spinner("Extracting OCR text from PDF images..."):
            extracted_texts = extract_text_from_images(images)

        # Create three output tabs.
        tabs = st.tabs(["Detected Content", "OCRmyPDF Output", "Semantic Chunks"])

        with tabs[0]:
            st.subheader("Detected Content with Red Bounding Boxes")
            with st.expander("Preview Detected Content"):
                for idx, img in enumerate(images, start=1):
                    preview_img = draw_bounding_boxes_on_image(img.copy())
                    st.image(preview_img, caption=f"Page {idx}", use_column_width=True)

            export_format = st.radio("Select Export Format", ("Excel", "Word", "Tables (Tesseract)"), key="export_format")
            if export_format == "Tables (Tesseract)":
                with st.spinner("Extracting table data via Tesseract TSV output..."):
                    table_dfs = extract_tables_from_images_with_tesseract(images)
                if table_dfs:
                    st.success(f"Extracted table data from {len(table_dfs)} page(s).")
                    for idx, table_df in enumerate(table_dfs, start=1):
                        st.write(f"### Page {idx} TSV Data")
                        st.dataframe(table_df)
                        csv_data = table_df.to_csv(index=False).encode("utf-8")
                        st.download_button(
                            label=f"Download Page {idx} Data as CSV",
                            data=csv_data,
                            file_name=f"page_{idx}_table_data.csv",
                            mime="text/csv"
                        )
                else:
                    st.error("No table-like data could be extracted using Tesseract TSV output.")
            else:
                if extracted_texts:
                    st.success("OCR text extraction complete.")
                    if export_format == "Excel":
                        excel_file = create_excel_file(extracted_texts)
                        st.download_button(
                            label="Download Excel File",
                            data=excel_file,
                            file_name="extracted_text.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                        )
                    elif export_format == "Word":
                        word_file = create_word_file(extracted_texts)
                        st.download_button(
                            label="Download Word Document",
                            data=word_file,
                            file_name="extracted_text.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                else:
                    st.error("No OCR text was extracted from the PDF images.")

        with tabs[1]:
            st.subheader("OCRmyPDF Processed PDF (Searchable)")
            if ocr_pdf is not None:
                st.success("OCRmyPDF processing complete.")
                st.download_button(
                    label="Download OCRmyPDF Processed PDF",
                    data=ocr_pdf,
                    file_name="ocr_processed.pdf",
                    mime="application/pdf"
                )
                try:
                    # Embed the PDF using an <object> tag.
                    b64_pdf = base64.b64encode(ocr_pdf).decode("utf-8")
                    pdf_display = f'<object data="data:application/pdf;base64,{b64_pdf}" type="application/pdf" width="700" height="900"></object>'
                    st.markdown(pdf_display, unsafe_allow_html=True)
                except Exception as e:
                    st.info("PDF preview embedding not available; please download the file.")
            else:
                st.error("OCRmyPDF processing failed.")

        with tabs[2]:
            st.subheader("Semantic Chunks from OCR Text using LlamaIndex")
            if not extracted_texts or all(text.strip() == "" for text in extracted_texts):
                st.error("No OCR text available for semantic chunking.")
            else:
                full_text = "\n\n".join(extracted_texts)
                chunks = get_semantic_chunks(full_text, chunk_size=512, chunk_overlap=50)
                if chunks:
                    st.success(f"Extracted {len(chunks)} semantic chunks.")
                    for i, chunk in enumerate(chunks, start=1):
                        token_count = len(chunk.split())
                        st.write(f"**Chunk {i}** (approx. {token_count} tokens):")
                        st.text_area("", chunk, height=150)
                else:
                    st.error("No semantic chunks could be extracted.")
